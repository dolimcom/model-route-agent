from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"D:\project\路由组件设计规划.docx")

NAVY = "17324D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "287271"
LIGHT_BLUE = "E8EEF5"
LIGHT_TEAL = "E8F3F1"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
DARK = "17212B"
WHITE = "FFFFFF"
GOLD = "8A6500"
RED = "9B1C1C"


def set_run_font(run, ascii_font="Calibri", east_asia="Microsoft YaHei", size=None,
                 color=None, bold=None, italic=None):
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def paragraph_border_bottom(paragraph, color=BLUE, size=12, space=4):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MID_GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, value, fld_end])
    end = paragraph.add_run(" 页")
    set_run_font(end, size=9, color=MID_GRAY)


def add_numbering(doc):
    numbering = doc.part.numbering_part.element

    def create_abstract(abstract_id, num_fmt, text, left=540, hanging=260, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.extend([tabs, ind])
        lvl.extend([start, fmt, lvl_text, suff, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), font)
            fonts.set(qn("w:hAnsi"), font)
            r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    def create_num(num_id, abstract_id):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    create_abstract(101, "bullet", "•", font="Arial")
    create_num(101, 101)
    create_abstract(102, "decimal", "%1.")
    create_num(102, 102)
    return 101, 102


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.45)
section.footer_distance = Inches(0.45)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.22

for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, DARK_BLUE, 10, 5)):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

caption = styles["Caption"]
caption.font.name = "Calibri"
caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
caption.font.size = Pt(9)
caption.font.color.rgb = RGBColor.from_string(MID_GRAY)

bullet_num, decimal_num = add_numbering(doc)

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = hp.add_run("MODEL-ROUTE SEMANTIC ROUTER   |   组件设计规划")
set_run_font(hr, size=8.5, color=MID_GRAY, bold=True)
paragraph_border_bottom(hp, color="D8DEE5", size=5, space=3)

footer = section.footer
fp = footer.paragraphs[0]
add_page_field(fp)


def add_para(text="", bold_prefix=None, italic=False, color=DARK, size=10.5,
             align=None, after=6, keep=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.22
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=size, color=color, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=size, color=color, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, italic=italic)
    return p


def add_bullets(items, level=0):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.22
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), str(level))
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), str(bullet_num))
        num_pr.extend([ilvl, num_id])
        p_pr.append(num_pr)
        r = p.add_run(item)
        set_run_font(r, size=10.3)


def add_numbered(items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.22
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), str(decimal_num))
        num_pr.extend([ilvl, num_id])
        p_pr.append(num_pr)
        r = p.add_run(item)
        set_run_font(r, size=10.3)


def add_callout(title, text, fill=LIGHT_BLUE, accent=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.18
    paragraph_shading(p, fill)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), accent)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    r1 = p.add_run(title + "  ")
    set_run_font(r1, size=10.3, color=accent, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.3, color=DARK)


def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.08
    paragraph_shading(p, "F6F8FA")
    r = p.add_run(text)
    set_run_font(r, ascii_font="Consolas", east_asia="Microsoft YaHei", size=8.7, color="243447")


def add_table(headers, rows, widths_dxa, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    set_repeat_header(table.rows[0])
    for idx, header_text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header_text)
        set_run_font(r, size=9.3, color=NAVY, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            r = p.add_run(str(value))
            set_run_font(r, size=9.1)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def page_break():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# Cover
add_para("TECHNICAL DESIGN SPECIFICATION", color=TEAL, size=9.5, after=12)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run("路由组件设计规划")
set_run_font(r, size=27, color=NAVY, bold=True)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(18)
r = p.add_run("Model-Route Semantic Router")
set_run_font(r, size=17, color=BLUE, bold=True)

meta = add_table(
    ["项目", "内容"],
    [
        ("交付形态", "独立 Java 语义路由组件 + Spring Boot Starter"),
        ("核心能力", "Embedding 语义匹配、边界决策、降级、人工干预、评测与可观测性"),
        ("本地集成", "Ollama / LM Studio / LocalAI 适配与能力探测"),
        ("编制日期", date.today().isoformat()),
        ("目标版本", "0.1.0-SNAPSHOT"),
    ],
    [1900, 7460],
    header_fill=LIGHT_TEAL,
)
add_callout(
    "设计结论",
    "路由质量优先于功能数量。组件必须能够解释为什么选择某条路由、何时拒绝选择、如何被人工纠正，并用独立数据集证明效果；否则不发布为可复用 Starter。",
    fill="EDF5FA",
    accent=TEAL,
)
add_para("适用读者：组件开发者、主 Agent 项目维护者、开源使用者与技术面试评审。", color=MID_GRAY, size=9.5)

page_break()

# Executive summary
doc.add_heading("1. 目标与设计原则", level=1)
add_para("本组件在业界可归类为 Semantic Router、LLM Router 或 Model Routing Middleware。它接收用户文本及可选对话上下文，输出可解释的路由目标、置信度、候选分数、降级原因和配置版本；模型调用由上层应用或可选编排适配器完成。")
add_table(
    ["原则", "落地要求"],
    [
        ("质量可证明", "提供隔离评测集、基线、混淆矩阵、fallback 指标和错误案例。"),
        ("失败优于误路由", "低分、近似平局、未知语义和编码失败均进入可配置降级链。"),
        ("配置优于改代码", "新增 route、样例、阈值、权重和目标模型只修改声明式配置。"),
        ("核心与框架解耦", "core 不依赖 Spring、数据库、主 Agent 的 TaskType 或具体模型厂商。"),
        ("本地优先", "Embedding、发现、决策和 LLM 推理可全部在本机完成。"),
        ("默认保护隐私", "日志默认记录输入哈希和摘要，不默认保存原始用户文本。"),
    ],
    [1900, 7460],
)

doc.add_heading("1.1 范围", level=2)
add_bullets([
    "实现可插拔 SemanticEncoder、RouteIndex、SimilarityScorer、RoutingPolicy 和 RouteDefinitionProvider。",
    "支持语义相似度、规则预过滤、加权混合、人工强制和 fallback 链等策略组合。",
    "支持配置热加载、不可变路由快照、版本化决策和 last-known-good 回退。",
    "提供本地模型发现适配、Actuator 端点、结构化日志、Micrometer 指标、CLI 评测和示例应用。",
    "发布 Spring Boot Starter，并允许纯 Java 项目只依赖 core。",
])

doc.add_heading("1.2 非目标", level=2)
add_bullets([
    "0.1 版本不引入 Milvus、Elasticsearch 或分布式向量数据库；路由样本使用内存索引。",
    "不负责长期对话存储、用户鉴权、计费或业务模型的完整生命周期管理。",
    "不承诺单一阈值适合所有领域；阈值必须通过使用者自己的验证集校准。",
    "不自动扫描局域网或任意地址；本地发现只探测显式配置或受限 loopback 端点。",
])

doc.add_heading("1.3 核心请求链", level=2)
add_code(
    "RoutingRequest\n"
    "  -> InputNormalizer\n"
    "  -> RouteConfigSnapshot(version)\n"
    "  -> SemanticEncoder\n"
    "  -> InMemoryRouteIndex\n"
    "  -> SimilarityScorer + RuleSignals\n"
    "  -> RoutingPolicy(top1, top2, threshold, margin)\n"
    "  -> ManualOverride / FallbackChain\n"
    "  -> RoutingResult + RouteTrace + Metrics"
)

page_break()

# Quality
doc.add_heading("2. 路由效果质量设计", level=1)
add_callout("核心验收", "任何正式发布必须同时给出测试集版本、路由配置版本、Embedding 模型版本和评测报告。只展示几个成功 Demo 不视为质量证明。", fill="FFF8E8", accent=GOLD)

doc.add_heading("2.1 路由分数与候选聚合", level=2)
add_para("每个 route 由 routeId、description、utterances、target 与 metadata 构成。输入向量与 route 样例向量计算余弦相似度，单条 route 的聚合方式可配置，而不是固定使用最大值。")
add_table(
    ["聚合策略", "说明", "适用场景"],
    [
        ("MAX", "取最相似样例；召回高但容易被单个异常样例影响。", "样例少、意图清晰的 MVP。"),
        ("TOP_K_MEAN", "取前 K 个相似度均值，降低单点噪声。", "默认推荐，K=3。"),
        ("CENTROID", "输入与 route 中心向量比较，速度快。", "样例分布紧凑、规模较大。"),
        ("WEIGHTED_HYBRID", "语义分、规则分、上下文分加权。", "需要兼顾强规则和隐式语义。"),
    ],
    [1700, 3900, 3760],
)
add_code(
    "semanticScore = aggregate(cosine(inputVector, routeExampleVectors))\n"
    "finalScore    = wSemantic * semanticScore + wRule * ruleScore + wContext * contextScore\n"
    "margin        = top1.finalScore - top2.finalScore"
)

doc.add_heading("2.2 边界情况与确定性决策", level=2)
add_table(
    ["场景", "判定", "默认行为", "可配置项"],
    [
        ("不在任何预设范围", "top1 < minScore", "返回 FALLBACK_OUT_OF_SCOPE", "min-score、fallback-route"),
        ("两个目标非常接近", "margin < minMargin", "返回 FALLBACK_AMBIGUOUS", "min-margin、tie-policy"),
        ("置信度低", "confidence < minConfidence", "降级到 LLM classifier 或默认模型", "fallback-chain"),
        ("完全同分", "top1 == top2", "确定性 fallback，不按集合顺序选择", "tie-policy"),
        ("空输入", "blank / only punctuation", "INVALID_INPUT", "reject-blank"),
        ("向量异常", "维度不符、零向量、NaN", "ENCODER_ERROR 或 last-known-good", "encoder-failure-policy"),
        ("配置重载失败", "校验或预编码失败", "继续使用上一有效快照", "reload-failure-policy"),
    ],
    [1850, 1900, 3050, 2560],
)

doc.add_heading("2.3 置信度", level=2)
add_para("原始余弦相似度不直接等价于概率。0.1 版本返回 rawScore、margin 与归一化 confidence，并在评测报告中单独披露；后续可通过验证集使用逻辑回归或温度缩放校准。禁止把未经校准的 0.82 描述成 82% 正确率。")
add_bullets([
    "RoutingResult 同时保留 topCandidates，默认返回前 3 个候选。",
    "confidence 计算器通过 ConfidenceCalibrator SPI 可替换。",
    "阈值按全局默认值设置，也允许 route 级覆盖。",
    "输出 decisionReason，例如 LOW_TOP_SCORE、SMALL_MARGIN、MANUAL_OVERRIDE、ENCODER_UNAVAILABLE。",
])

doc.add_heading("2.4 人工干预与纠正", level=2)
add_table(
    ["能力", "接口/配置", "行为"],
    [
        ("请求级提示", "RoutingRequest.routeHint", "SOFT 模式加权；HARD 模式强制但记录 override。"),
        ("静态覆盖", "overrides.yml", "按精确短语、正则或业务标签覆盖，支持优先级和过期时间。"),
        ("结果纠正", "POST /feedback 或 FeedbackRecorder", "保存 predicted、corrected、scores、版本和备注。"),
        ("闭环更新", "evaluation promote-feedback", "经人工审核后将纠正样本加入 route utterances，禁止自动污染生产样本。"),
    ],
    [1700, 3000, 4660],
)

page_break()

# Evaluation
doc.add_heading("3. 评测体系与质量门禁", level=1)
doc.add_heading("3.1 数据集设计", level=2)
add_para("路由样例和测试集必须物理隔离，避免把 route utterances 原样放入测试集造成数据泄漏。仓库提供 datasets/routes 与 datasets/evaluation 两套目录，并为每次评测生成 datasetVersion。")
add_table(
    ["集合", "建议规模", "内容"],
    [
        ("Route utterances", "每 route 15-30 条", "用于建立索引；覆盖直白、口语、隐式表达和中英文变体。"),
        ("In-domain evaluation", "每 route 至少 30 条", "与 utterances 不重复的真实风格问题。"),
        ("OUT_OF_SCOPE", "至少 45 条", "不属于任何 route 的输入。"),
        ("AMBIGUOUS", "至少 45 条", "两个或多个 route 都合理的边界输入。"),
        ("Robustness", "至少 30 条", "空白、超长、错别字、混合语言、代码与自然语言混排。"),
    ],
    [2100, 1700, 5560],
)
add_code(
    '{"id":"eval-001","text":"这段东西一运行就停了",'
    '"expectedRoute":"coding","expectedFallback":false,'
    '"tags":["implicit","zh"],"source":"handcrafted","notes":""}'
)

doc.add_heading("3.2 指标", level=2)
add_table(
    ["指标", "定义与用途"],
    [
        ("Accuracy / Macro-F1", "衡量各 route 分类效果，Macro-F1 防止大类掩盖小类。"),
        ("Selective accuracy", "只统计被系统接受的结果，衡量阈值后的可靠性。"),
        ("Coverage", "非 fallback 样本占比；与 selective accuracy 联合观察。"),
        ("Fallback Precision/Recall/F1", "判断系统能否正确拒绝 OOD 和 ambiguous 输入。"),
        ("Confusion matrix", "展示常见混淆方向，例如 literary 与 daily。"),
        ("Calibration / ECE", "检查 confidence 是否可信。"),
        ("Latency p50/p95", "区分 encode、score、policy 和 totalLatency。"),
    ],
    [2500, 6860],
)

doc.add_heading("3.3 对比基线", level=2)
add_bullets([
    "Random baseline：在可选 route 中固定种子随机选择。",
    "Fixed baseline：永远选择训练集中占比最大的 route。",
    "Keyword baseline：复用主 Agent 的规则关键词路由。",
    "Semantic candidate：本组件语义路由；评测时保持相同测试集。",
    "Ablation：分别关闭 margin、规则分和上下文分，解释每个机制的贡献。",
])

doc.add_heading("3.4 错误案例报告", level=2)
add_para("评测命令必须输出 errors.jsonl 与 Markdown 摘要。每条错误至少包含输入、期望 route、实际 route、topCandidates、rawScore、margin、fallbackReason、配置版本和初步归因。")
add_table(
    ["归因标签", "示例原因", "处理方式"],
    [
        ("DATA_GAP", "route 样例未覆盖口语表达", "补充经审核 utterance。"),
        ("LABEL_AMBIGUITY", "人工标签本身存在争议", "改为 ambiguous 或多标签。"),
        ("ENCODER_LIMIT", "Embedding 无法区分领域术语", "替换编码器或领域微调。"),
        ("THRESHOLD", "正确候选第一但被阈值拒绝", "使用验证集重新校准。"),
        ("CONTEXT_MISSING", "短句依赖上一轮对话", "启用 contextComposer。"),
    ],
    [2200, 4300, 2860],
)

doc.add_heading("3.5 0.1 发布门禁", level=2)
add_bullets([
    "Macro-F1 >= 0.80；各核心 route F1 不低于 0.70。",
    "Selective accuracy >= 0.90，同时 coverage >= 0.75。",
    "Fallback F1 >= 0.80，必须覆盖 OOD 与 ambiguous。",
    "随机、固定和关键词基线全部运行，语义路由至少优于固定基线 15 个百分点。",
    "核心评分 p95 < 10 ms；本地 Embedding 端到端 p95 目标 < 1 s，并披露硬件。",
    "所有错误案例可导出，评测命令使用固定随机种子且结果可复现。",
])

page_break()

# Strategies/config
doc.add_heading("4. 可配置路由策略", level=1)
doc.add_heading("4.1 策略组合", level=2)
add_table(
    ["策略", "职责", "是否默认"],
    [
        ("SemanticSimilarityStrategy", "Embedding + 向量索引 + 聚合评分。", "是"),
        ("RuleSignalStrategy", "正则、关键词、代码块等强信号；只产出分数，不直接硬编码返回。", "可选"),
        ("ContextAwareStrategy", "将最近对话摘要或当前文件摘要合入编码输入。", "可选"),
        ("ManualOverrideStrategy", "处理 HARD/SOFT hint 和静态 override。", "可选，优先级最高"),
        ("FallbackStrategy", "按原因执行 DEFAULT_ROUTE、LLM_CLASSIFIER、REJECT 等降级。", "是"),
    ],
    [2600, 5200, 1560],
)

doc.add_heading("4.2 声明式配置示例", level=2)
add_code(
    "semantic-router:\n"
    "  enabled: true\n"
    "  encoder:\n"
    "    type: ollama\n"
    "    base-url: http://localhost:11434\n"
    "    model: bge-m3\n"
    "    timeout: 3s\n"
    "  scoring:\n"
    "    aggregation: top-k-mean\n"
    "    top-k: 3\n"
    "    min-score: 0.62\n"
    "    min-margin: 0.08\n"
    "    min-confidence: 0.70\n"
    "  fallback:\n"
    "    chain: [llm-classifier, default-route]\n"
    "    default-route: general\n"
    "  reload:\n"
    "    enabled: true\n"
    "    debounce: 500ms\n"
    "  routes:\n"
    "    - id: coding\n"
    "      target-model: local-coder\n"
    "      utterances: [\"这段程序为什么突然停了\", \"帮我定位运行异常\"]\n"
    "      min-score: 0.66\n"
    "    - id: literary\n"
    "      target-model: local-writer\n"
    "      utterances: [\"让这段文字更有画面感\", \"把语气写得更克制\"]"
)

doc.add_heading("4.3 热加载", level=2)
add_numbered([
    "监听显式配置文件变化，进行 500 ms debounce，避免编辑器连续写入导致重复加载。",
    "解析到临时 RouteConfig，执行 schema、routeId 唯一性、样例非空、阈值范围和目标引用校验。",
    "批量编码新增或变化样例，构造不可变 RouteSnapshot，并计算 version/hash。",
    "通过 AtomicReference 一次性切换快照；正在执行的请求继续使用旧快照。",
    "任何阶段失败都保留 last-known-good，并记录 reload.failure 指标与结构化错误。",
])
add_callout("新增 route", "正常情况下只需增加配置和样例，不修改 Java 代码。只有新增评分算法、编码协议或降级类型时才实现新的 SPI。", fill=LIGHT_TEAL, accent=TEAL)

page_break()

# Local integration
doc.add_heading("5. 本地模型集成", level=1)
doc.add_heading("5.1 适配层", level=2)
add_para("组件通过 LocalModelProvider SPI 统一模型发现、Embedding 与可选 Chat Completion。核心路由只依赖 SemanticEncoder；完整本地链路由 sample/orchestrator 模块演示，避免 core 强耦合模型服务。")
add_table(
    ["本地运行时", "发现", "Embedding", "Chat", "适配说明"],
    [
        ("Ollama", "GET /api/tags", "POST /api/embed", "Native 或 /v1/chat/completions", "优先原生 Embedding；支持 OpenAI 兼容 Chat。"),
        ("LM Studio", "GET /v1/models", "POST /v1/embeddings", "POST /v1/chat/completions", "使用 OpenAI-compatible adapter。"),
        ("LocalAI", "能力探测 /v1/models", "OpenAI-compatible adapter", "OpenAI-compatible adapter", "端点以集成测试为准，不在 core 写死厂商假设。"),
    ],
    [1500, 1650, 1900, 2100, 2210],
)
add_para("核对依据：Ollama 官方 OpenAI compatibility 与 List models 文档、LM Studio 官方 OpenAI Compatibility Endpoints、Spring Boot 官方 Creating Your Own Auto-configuration。LocalAI 官网访问不稳定，采用官方 OpenAI-compatible 定位设计，并将端点契约纳入集成测试。", color=MID_GRAY, size=8.8)

doc.add_heading("5.2 自动检测与环境适配", level=2)
add_bullets([
    "显式配置优先；只有 discovery.enabled=true 时才探测 loopback。",
    "默认候选仅包含 localhost:11434（Ollama）与 localhost:1234（LM Studio）；LocalAI 端口必须显式配置，避免与主应用 8080 冲突。",
    "探测超时默认 300 ms，失败不阻塞 Spring 启动；检测结果带 provider、endpoint、models、capabilities 和 checkedAt。",
    "不读取或猜测模型文件系统路径，通过服务 API 获取模型清单；环境变量可覆盖 base URL、model 和 timeout。",
    "模型不存在时给出可执行诊断，例如 Missing embedding model: bge-m3，而不是空指针或连接拒绝堆栈。",
    "自动探测只允许 loopback allowlist，禁止将用户输入拼接成探测 URL，降低 SSRF 风险。",
])

doc.add_heading("5.3 完整本地链路", level=2)
add_code(
    "Local runtime starts\n"
    "  -> ModelDiscoveryClient detects embedding/chat models\n"
    "  -> SemanticEncoder encodes request locally\n"
    "  -> SemanticRouter selects route + targetModel\n"
    "  -> RouteTargetResolver validates available model\n"
    "  -> LocalChatClient invokes local LLM\n"
    "  -> RouteTrace records decision and inference latency"
)
add_callout("离线验收", "断开公网后，Ollama/LM Studio 已下载模型的情况下，示例应用仍能完成从输入、Embedding、路由到本地 LLM 回答的完整链路。", fill="EDF5FA", accent=TEAL)

page_break()

# Architecture
doc.add_heading("6. 系统设计与模块边界", level=1)
doc.add_heading("6.1 Maven 模块", level=2)
add_table(
    ["模块", "依赖边界", "交付内容"],
    [
        ("semantic-router-core", "纯 Java；不依赖 Spring/Web/数据库", "API、索引、评分、策略、快照与异常。"),
        ("semantic-router-spring-boot-autoconfigure", "Spring Boot autoconfigure", "Properties、Conditional beans、Actuator、reload。"),
        ("semantic-router-spring-boot-starter", "依赖 core + autoconfigure", "供业务项目一行 Maven 引入。"),
        ("semantic-router-evaluation", "CLI 与测试工具", "数据加载、基线、指标、混淆矩阵和错误报告。"),
        ("semantic-router-sample", "Spring Web + 本地模型 adapter", "可直接运行的 Dashboard 与端到端示例。"),
    ],
    [2900, 2900, 3560],
)

doc.add_heading("6.2 核心 SPI", level=2)
add_code(
    "SemanticRouter.route(RoutingRequest) -> RoutingResult\n"
    "SemanticEncoder.encode(List<String>) -> List<Vector>\n"
    "RouteDefinitionProvider.load() -> RouteConfig\n"
    "RouteIndex.search(Vector, limit) -> List<RouteScore>\n"
    "SimilarityScorer.score(input, route) -> double\n"
    "RoutingPolicy.decide(candidates, context) -> RouteDecision\n"
    "ConfidenceCalibrator.calibrate(features) -> double\n"
    "FeedbackRecorder.record(RouteFeedback)\n"
    "ModelDiscoveryClient.discover() -> LocalModelCatalog"
)

doc.add_heading("6.3 RoutingResult 契约", level=2)
add_table(
    ["字段", "说明"],
    [
        ("routeId / target", "最终 route 与可选目标模型。"),
        ("status", "ROUTED、FALLBACK、REJECTED、ERROR。"),
        ("confidence", "经校准或显式标注为未校准的置信度。"),
        ("topCandidates", "候选 route、semantic/rule/context/final 分数。"),
        ("reasonCode / explanation", "机器可读原因与人类可读解释。"),
        ("manualOverride", "是否人工干预、模式和来源。"),
        ("configVersion / encoderVersion", "支持复盘和可复现评测。"),
        ("timings", "encode、search、policy、total 毫秒。"),
        ("traceId", "关联日志、指标、反馈和上层模型调用。"),
    ],
    [2800, 6560],
)

doc.add_heading("6.4 并发与可靠性", level=2)
add_bullets([
    "RouteSnapshot、索引和 route 定义不可变；请求线程只读，无全局写锁。",
    "热加载在后台构建新快照，原子切换；失败不破坏现有服务。",
    "Embedding 支持批量、超时、有限重试和可选缓存；缓存 key 包含 encoderVersion。",
    "外部本地服务调用设置连接/响应超时、并发上限与熔断，不无限重试。",
    "所有异常映射到稳定错误码，Spring Starter 不把第三方堆栈直接暴露给 API。",
])

page_break()

# Observability
doc.add_heading("7. 可观测性与复盘", level=1)
doc.add_heading("7.1 结构化日志", level=2)
add_code(
    '{"event":"semantic_route_decision","traceId":"...",'
    '"inputHash":"sha256:...","routeId":"coding","status":"ROUTED",'
    '"topScore":0.81,"secondScore":0.66,"margin":0.15,'
    '"confidence":0.88,"reasonCode":"ACCEPTED",'
    '"configVersion":"routes-a13f","encoderVersion":"bge-m3",'
    '"encodeMs":42,"policyMs":1,"totalMs":45}'
)
add_bullets([
    "默认不记录原始输入；可在本地开发配置中 opt-in，并支持最大长度与敏感字段脱敏。",
    "记录 topCandidates 而非只记录最终结果，便于分析边界样本。",
    "反馈、配置重载、模型发现和 provider 错误使用独立 eventType。",
])

doc.add_heading("7.2 指标", level=2)
add_table(
    ["指标", "标签/用途"],
    [
        ("semantic_router_requests_total", "status、routeId、reasonCode"),
        ("semantic_router_fallback_total", "fallbackReason"),
        ("semantic_router_latency", "stage=encode/search/policy/total"),
        ("semantic_router_confidence", "routeId 分布"),
        ("semantic_router_reload_total", "result=success/failure"),
        ("semantic_router_provider_health", "provider、capability"),
        ("semantic_router_feedback_total", "predicted、corrected"),
    ],
    [3600, 5760],
)

doc.add_heading("7.3 本地 Dashboard", level=2)
add_para("Starter 提供只读 Actuator 端点；sample 模块提供本地 Web Dashboard。Dashboard 不进入 core，也不强制业务项目引入前端依赖。")
add_bullets([
    "最近路由记录：输入摘要、route、候选分数、置信度、reason 和耗时。",
    "配置状态：当前版本、加载时间、route 数、样例数、上次重载错误。",
    "Provider 状态：Ollama/LM Studio/LocalAI 可用性与模型列表。",
    "评测页面：指标、混淆矩阵、错误案例和基线对比。",
    "人工纠正：选择正确 route、填写备注并导出 feedback.jsonl。",
])
add_callout("端点建议", "GET /actuator/semanticrouter、GET /actuator/semanticrouter/routes、POST /api/feedback（仅 sample 或上层应用启用）。生产默认关闭详细输入展示。", fill=LIGHT_TEAL, accent=TEAL)

page_break()

# OSS/quickstart
doc.add_heading("8. 开箱即用与开源质量", level=1)
doc.add_heading("8.1 Clone-to-run 路径", level=2)
add_code(
    "git clone <repository>\n"
    "cd model-route-semantic-router\n"
    "ollama pull bge-m3\n"
    "mvn clean verify\n"
    "mvn -pl semantic-router-sample spring-boot:run\n"
    "# open http://localhost:8080"
)
add_para("默认 sample 配置包含 coding、literary、math、daily、general 五类 route。未安装 Ollama 时，启动诊断明确提示安装/拉取命令；测试仍使用 DeterministicTestEncoder，不依赖网络。")

doc.add_heading("8.2 仓库必备内容", level=2)
add_table(
    ["文件/目录", "要求"],
    [
        ("README.md", "30 秒理解项目、架构图、三步 Quickstart、配置、评测结果、限制。"),
        ("docs/architecture.md", "模块、线程模型、热加载、失败策略和扩展点。"),
        ("docs/evaluation.md", "数据集、指标、硬件、基线、结果与错误案例。"),
        ("examples/", "纯 Java、Spring Boot、Ollama 三个最小示例。"),
        ("datasets/", "routes 与 evaluation 分离，携带 license/source。"),
        ("src/test", "核心单测、自动装配测试、契约测试、配置重载和边界测试。"),
        ("LICENSE / CONTRIBUTING", "建议 Apache-2.0；说明提交、测试和行为规范。"),
        ("CHANGELOG.md", "版本能力、兼容性和 breaking changes。"),
    ],
    [2600, 6760],
)

doc.add_heading("8.3 测试矩阵", level=2)
add_bullets([
    "单元测试：余弦相似度、聚合、阈值、margin、tie、fallback、人工覆盖。",
    "属性测试：随机向量下分数范围、对称性、零向量与 NaN。",
    "配置测试：重复 route、空 utterance、非法阈值、目标缺失、热加载失败。",
    "Spring 测试：ApplicationContextRunner 验证 ConditionalOnMissingBean 和属性绑定。",
    "契约测试：Ollama/LM Studio/LocalAI adapter 使用本地 mock server 固定响应。",
    "端到端测试：可选 profile 连接真实本地运行时，不进入默认 CI。",
    "回归测试：固定评测集与最低指标门禁，PR 下降超过阈值则失败。",
])

page_break()

# Security and roadmap
doc.add_heading("9. 安全、隐私与运维约束", level=1)
add_table(
    ["风险", "控制"],
    [
        ("日志泄露用户输入", "默认 hash/摘要；原文记录显式 opt-in；支持脱敏。"),
        ("配置热加载污染", "临时解析、全量校验、原子切换、last-known-good。"),
        ("自动发现造成 SSRF", "只允许配置与 loopback allowlist；短超时；禁止用户输入 URL。"),
        ("Embedding 服务卡死", "连接/读取超时、并发上限、熔断与明确降级。"),
        ("反馈数据污染", "纠正记录与生产样例隔离，人工审核后再 promote。"),
        ("不可复现路由", "记录 configVersion、encoderVersion、datasetVersion 和随机种子。"),
    ],
    [2600, 6760],
)

doc.add_heading("10. 实施顺序", level=1)
add_table(
    ["阶段", "交付", "完成定义"],
    [
        ("A. Core", "API、向量、索引、评分、policy、fallback", "边界单测全部通过；无 Spring 依赖。"),
        ("B. Config", "YAML、校验、快照、热加载", "错误配置不替换有效快照。"),
        ("C. Local", "Ollama encoder、发现、LM Studio/LocalAI adapter", "至少 Ollama 完成本地真实链路。"),
        ("D. Starter", "自动装配、Actuator、配置元数据", "主项目只添加依赖和 YAML 即可使用。"),
        ("E. Evaluation", "数据集、基线、指标、错误报告", "满足 0.1 质量门禁或明确披露未达项。"),
        ("F. OSS", "README、sample、Dashboard、license", "新环境按 Quickstart 可运行。"),
    ],
    [1700, 4200, 3460],
)

doc.add_heading("11. 最终验收清单", level=1)
add_bullets([
    "低分、近似平局、同分、OOD、空输入和 encoder 故障都有确定性结果。",
    "新增 route 只改配置；新增算法通过 SPI，不修改核心编排。",
    "路由配置热加载失败时服务继续使用上一版本。",
    "支持 SOFT/HARD hint、静态 override 和人工 feedback。",
    "评测报告包含数据集说明、基线、Macro-F1、coverage、fallback F1、混淆矩阵和错误案例。",
    "至少 Ollama 完成离线端到端；LM Studio 与 LocalAI 通过 adapter 契约测试。",
    "日志、Actuator 和 Dashboard 能复盘 route、候选分数、原因、版本与耗时。",
    "git clone 后按 README 在 10 分钟内完成测试和 sample 启动。",
    "Starter 可被 model-route-agent 通过 Maven 依赖独立引入，不反向依赖主项目。",
])

doc.add_heading("12. 风险与取舍", level=1)
add_table(
    ["风险/取舍", "决策"],
    [
        ("本地 Embedding 硬件差异", "报告硬件与延迟；编码器可替换；默认批量和缓存。"),
        ("阈值过拟合小数据集", "单独 validation/test；披露样本量；提供用户校准命令。"),
        ("自动发现复杂度", "仅做受限 loopback 探测，显式配置始终优先。"),
        ("Dashboard 扩大 Starter 体积", "Dashboard 仅在 sample，Starter 提供 Actuator 数据。"),
        ("多策略难解释", "分项输出 semantic/rule/context/final 分数与 reasonCode。"),
        ("一天内范围过大", "先完成 Core + Ollama + Evaluation；LM Studio/LocalAI 以标准 adapter 和契约测试收敛。"),
    ],
    [3000, 6360],
)

doc.add_heading("13. 官方参考", level=1)
add_bullets([
    "Ollama OpenAI compatibility: https://docs.ollama.com/api/openai-compatibility",
    "Ollama List models: https://docs.ollama.com/api/tags",
    "LM Studio OpenAI Compatibility Endpoints: https://lmstudio.ai/docs/developer/openai-compat",
    "LocalAI official project: https://github.com/mudler/LocalAI",
    "Spring Boot - Creating Your Own Auto-configuration: https://docs.spring.io/spring-boot/reference/features/developing-auto-configuration.html",
])

# Keep headings with following content and prevent widows where possible.
for paragraph in doc.paragraphs:
    if paragraph.style.name.startswith("Heading"):
        paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.widow_control = True

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
